import seaborn as sns 
import streamlit as st
import matplotlib.pyplot as plt
import pandas as pd
import numpy as np
import os
import io
import xlsxwriter
from sklearn.linear_model import LinearRegression
from sklearn.metrics import r2_score
from docx import Document
from functools import wraps

st.set_page_config(page_icon= 'jabuti-05.png', page_title= 'Facilitador')

class CacheDataAPI:
    def __init__(self):
        self._cached_data = {}

    def __call__(self, func):
        @wraps(func)
        def wrapped_func(*args, **kwargs):
            cache_key = hash((func, args, frozenset(kwargs.items())))
            if cache_key not in self._cached_data:
                self._cached_data[cache_key] = func(*args, **kwargs)
            return self._cached_data[cache_key]
        return wrapped_func

    def clear(self):
        self._cached_data = {}

# Cria uma instância da API de cache de dados
st.cache_data = CacheDataAPI()




st.title('Bem-vindo ao tratamento de ICP!')
st.write('Software desenvolvido por Miguel O.')
st.write('')
icp_file = st.file_uploader(label = 'Selecione o arquivo Excel com os dados ICP:', type=['xlsx'])


if icp_file:
    icp = pd.read_excel(icp_file)
    csv_buffer = io.StringIO()
    icp.to_csv(csv_buffer, index=False)
    csv_buffer.seek(0)
    grafico = pd.read_csv(csv_buffer)

    word_file = st.file_uploader(label= 'Selecione o arquivo Word com os comprimentos de onda:', type=['docx'])

    def find_highlighted_cells(doc):
        highlighted_cells = []
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():  # Verifica se a célula não está vazia
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                if run.font.highlight_color is not None:
                                    highlighted_cells.append(cell.text)
        return highlighted_cells

    if word_file:
        doc = Document(word_file)

        highlighted_cells = find_highlighted_cells(doc)

        df_ana = pd.DataFrame(highlighted_cells)
        
        dataframes_Color=[]
        for nome_coluna, dados_coluna in df_ana.items():
            novo_df = pd.DataFrame({nome_coluna: dados_coluna})
            novo_df = novo_df.dropna()
            dataframes_Color.append(novo_df)
        for i, df in enumerate(dataframes_Color):
            dataframes_Color[i] = df.reset_index(drop=True)
        df_ana = pd.concat(dataframes_Color, axis=1)
        df_ana = df_ana.replace(',', '.', regex=True)


        dados_tabelas = []
        # Aqui eu mexo no word
        marcacoes = []
        for tabela in doc.tables:
            dados_tabela = []
            for linha in tabela.rows:
                linha_tabela = [celula.text for celula in linha.cells]
                dados_tabela.append(linha_tabela)
                if any('marca texto' in celula.text for celula in linha.cells):
                    linha_tabela = [celula.text for celula in linha.cells]
                    dados_tabela.append(linha_tabela)
            dados_tabelas.append(pd.DataFrame(dados_tabela[1:], columns=dados_tabela[0]))

        df_Word = pd.concat(dados_tabelas, ignore_index=True)
        df = df_Word.drop(columns= ['Ordem'])


        dataframes_modificados = []

        for nome_coluna, dados_coluna in df.items():
            novo_df = pd.DataFrame({nome_coluna: dados_coluna})
            novo_df = novo_df.dropna()
            dataframes_modificados.append(novo_df)
        for i, df in enumerate(dataframes_modificados):
            dataframes_modificados[i] = df.reset_index(drop=True)
        df_final = pd.concat(dataframes_modificados, axis=1)
        df_final = df_final.replace(',', '.', regex=True)

        def color_numbers(val):
            color = 'color: green' if val in df_ana.values.flatten() else ''
            return color

        # Aplicar a função ao DataFrame
        styled_df = df_final.style.applymap(color_numbers)

        # Exibir o DataFrame estilizado
        styled_df

        coluna_escolhida = st.selectbox('Defina o elemento:', df_final.columns)
        onda_escolhida = st.selectbox('Defina o comprimento de onda:', df_final[coluna_escolhida])

        Analito = str(coluna_escolhida + ' ' + onda_escolhida)



        # A partir daqui no excel

        Branco = grafico.loc[grafico["Type"] == "Blk"]
        Padrao = grafico.loc[grafico['Type'] == 'Std']
        Amostra = grafico.loc[grafico['Type'] == 'Samp']


        PontoB = Branco.loc[Branco['Element'] == Analito]
        PontoP = Padrao.loc[Padrao['Element'] == Analito]
        PontoA = Amostra.loc[Amostra['Element'] == Analito]


        limpeza = 'Limpeza'
        PontoA= PontoA[PontoA['Solution Label'] != limpeza]


        PontoP['Int'] = (PontoP['Int']) - int(PontoB['Int'])
        PontoA['Int'] = (PontoA['Int']) - int(PontoB['Int'])

        coluna1,coluna2,coluna3 = st.columns([6,2,2])
        
        x_original = PontoP['Soln Conc'].values
        y_original = PontoP['Int'].values
        selected_rows = []
        with coluna2:
            for index, row in PontoP.iterrows():
                selected = st.checkbox(f"{row['Soln Conc']}", key=index)
                if selected:
                    selected_rows.append(index)

        PontoP = PontoP.drop(selected_rows)

        x = PontoP['Soln Conc'].values
        y = PontoP['Int'].values
        model = LinearRegression().fit(x.reshape(-1, 1), y)
        y_prev = model.predict(x.reshape(-1, 1))
        r_quadrado = r2_score(y, y_prev)

        fig = plt.figure()

        left, width = .25, .5
        bottom, height = .25, .5
        right = left + width
        top = bottom + height

        plt.scatter(x_original, y_original, alpha=0)
        plt.scatter(x,y, color ='Black')
        fig.text(right, top, f'$R^2 = {r_quadrado:.5f}$', horizontalalignment='right', verticalalignment='bottom')
        plt.plot(x, y_prev, color='red')

        with coluna1:
            st.pyplot(plt)

        a = model.coef_[0]
        b = model.intercept_

        valor_referencia = PontoP['Int'].min()
        FGX = PontoA['Int']
        valor_maximo = PontoP['Int'].max()

        df_valores_ref = pd.DataFrame({'Máximo': valor_maximo, 'Minimo': valor_referencia}, index=[0])

        with coluna3:
            st.dataframe(df_valores_ref, hide_index=True)


        def highlight_row(row):
            if row['Int'] < valor_referencia:
                return ['color: red'] * len(row)
            elif row['Int'] > valor_maximo:
                return ['color: red'] * len(row)
            else:
                return ['color: green'] *len(row)

        PosiçãoNaCurva = []
        for index, row in PontoA.iterrows():
            if row['Int'] < valor_referencia:
                PosiçãoNaCurva.append('Abaixo')
            elif row['Int'] > valor_maximo:
                PosiçãoNaCurva.append('Acima')
            else:
                PosiçãoNaCurva.append('Dentro')

        PontoA['Int'] = (PontoA['Int']-b)/a

        PontoA = PontoA.rename(columns={'Solution Label':'Nome da Amostra', 'Int':'Concentração'})

        df_subset = PontoA[['Nome da Amostra', 'Concentração']]

        st.write('INSIRA OS FATORES DE DILUIÇÃO')

        # Copia df_subset para temp_df para manipulação
        temp_df = df_subset.copy()

        # Extrai o número antes do 'x' na coluna 'Nome da Amostra'
        temp_df['Fator de diluição'] = temp_df['Nome da Amostra'].str.extract(r'([\d.,]+)x', expand=False)

        # Remove pontos como separadores de milhar
        # Mantém os pontos decimais intactos
        def clean_fator_diluicao(value):
            # Se o valor é uma string, remova os pontos de milhar
            if isinstance(value, str):
                # Remove pontos de milhar e substitui vírgulas por pontos decimais
                value = value.replace('.', '')  # Remove todos os pontos
                value = value.replace(',', '.')  # Substitui vírgulas por pontos decimais
            try:
                return float(value)
            except ValueError:
                return 0  # Ou outra lógica para lidar com valores inválidos

        # Aplica a limpeza
        temp_df['Fator de diluição'] = temp_df['Fator de diluição'].apply(clean_fator_diluicao)

        # Cria o editor de dados no Streamlit
        df_subset = st.data_editor(temp_df[['Nome da Amostra', 'Fator de diluição']], height=300, hide_index=True)

        # Adiciona a coluna 'Concentração' de volta ao DataFrame editado
        df_subset['Concentração'] = PontoA['Concentração']

        # Realiza a multiplicação
        df_subset['Resultado'] = df_subset['Concentração'] * df_subset['Fator de diluição']
        df_subset['Resultado (ppm)'] = df_subset['Resultado'].round(2).astype(str)
        df_subset['Int'] = FGX
        df_subset['Posição na Curva'] = PosiçãoNaCurva
        df_subset['Int'] = df_subset['Int'].round(2).astype(int)

        styled_df = df_subset.style.apply(highlight_row, axis=1)
        df_subset['Fator de diluição'] = df_subset['Fator de diluição'].map(lambda x: f"{x:.1f}")

        st.table(styled_df)


        def convert_df(df_subset):
            output = io.BytesIO()
            with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
                df_subset.to_excel(writer, index=False, sheet_name='Sheet1')
            return output.getvalue()

        Excel = convert_df(df_subset)

        if Excel:
            st.download_button(
                label="Baixar Resultados",
                data=Excel,
                file_name='Resultados_ICP.xlsx',
                mime='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            )
    else:
        st.warning("Por favor, faça o upload do word.")
else:
 st.warning("Por favor, faça o upload do excel.")
